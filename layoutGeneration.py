#Project:ESA Procurement Ai - Phase 2 - Document Generation [Scope of Work]
#Author:Mohan Sundaram, Shakthi Gopalakrishnan, Vinay Gajendran
#Version:0.1

#---------------------------Import libraries-------------------#

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
import datetime

#-----------------------------Defintion------------------------------#

#-----------------------------Title section------------------------------#
under_text = """EMSTEEL 1.access to themerger ADX ticker of the listed entity was changed from ARKAN to EMSTEEL."""
print(len(under_text))
image = Image.open("utils\\assets\\image.png")

draw = ImageDraw.Draw(image)


if 300>len(under_text)>120:fontSize,linSize = 25 , 85
elif 300<len(under_text)<1000:fontSize,linSize = 15 , 140
else: fontSize,linSize=35,60
positionOffset = (520,130)

modified_under_text = ''

if len(under_text)>linSize:
    lineLength=0
    under_text_list = under_text.split(" ")
    len_under_text = len(under_text_list)
    for i in under_text_list:
        if lineLength<linSize and lineLength!=linSize:
            modified_under_text+=i+" "
            lineLength += len(i)
        else:
            modified_under_text+=f"\n{i} "
            lineLength =0
            positionOffset = (520,110)
        under_text = modified_under_text


font = ImageFont.truetype("arial.ttf", fontSize)

text_color = (0, 0, 0)

position = positionOffset #520-1750 x-axis

draw.text(position, under_text, fill=text_color, font=font)

overlay = Image.open("utils\\assets\\emsteel-logo-new.png")

overlay = overlay.resize((450, 100))
position = (40, 50)

image.paste(overlay, position, overlay)

image.save("utils\\assets\\image_.png")
print("Text drawn successfully!")

#-----------------------------------------------------Document content-------------------------------------#

# Create a new document
doc = Document()

doc.sections[0].top_margin=Pt(20)
doc.sections[0].left_margin=Pt(40)
doc.sections[0].right_margin=Pt(40)
doc.sections[0].bottom_margin=Pt(40)

#------------------------------------Document primary------------------------------------------#

prValue = "88742947578832948"
workTitle = " Annual Service Contract for 2025-2026"
typeJob = "Maintenance "
req = "RM1 Operations"
endUserFoculPoint = "Hussain Almoosawi/ Mohamed El Refaey"
date_ = datetime.datetime.now().date()
issue_no = ""
page_no = 1

#-------------------------Header section------------------------------------#
doc.add_picture("utils\\assets\\image_.png",width=Pt(530))
doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


#-------------------------------Table section----------------------------#
upper_table = doc.add_table(rows=5,cols=2)
upper_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
upper_table.autofit = False

#------------------------------Document Information----------------------#

upper_table.cell(0,0).merge(upper_table.cell(0,1))


para = upper_table.cell(0,0).paragraphs[0]
run = para.add_run("PR No.")
run.bold = True
para.add_run(f"{prValue}")
run.font.size = Pt(10)
# upper_table.cell(0,0).text += f"{prValue}"
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



# upper_table.cell(1,0).text = f"Work/Service Title: "
para = upper_table.cell(1,0).paragraphs[0]
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = para.add_run("Work/Service Title :")
run.bold = True
para.add_run(f" {workTitle}")
run.font.size = Pt(10)

# upper_table.cell(1,1).text = f"Type of Job :"
para = upper_table.cell(1,1).paragraphs[0]
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = para.add_run("Type of Job :")
run.bold = True
para.add_run(f" {typeJob}")
run.font.size = Pt(10)


# upper_table.cell(2,0).text = f"Name of Requisitioner:{req}"
para = upper_table.cell(2,0).paragraphs[0]
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = para.add_run("Name of Requisitioner:")
run.bold = True
para.add_run(f'{req}')
run.font.size = Pt(10)


# upper_table.cell(2,1).text = "Date:09/10/2025"
para=upper_table.cell(2,1).paragraphs[0]
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = para.add_run("Date:")
run.bold = True
para.add_run(str(date_))
run.font.size = Pt(10)


upper_table.cell(3,0).merge(upper_table.cell(3,1))

# upper_table.cell(3,0).text = f"End User Focal Point:{endUserFoculPoint}"
para = upper_table.cell(3,0).paragraphs[0]
para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = para.add_run("End User Focal Point:")
run.bold = True
para.add_run(f'{endUserFoculPoint}')
run.font.size = Pt(10)

upper_table.cell(4,0).merge(upper_table.cell(4,1))


#------------------------------Document context--------------------------------#
#------------------------------Document Description----------------------------#


desc = """\n  
1. Description* 
Elon Musk is a visionary entrepreneur and inventor. He was born in Pretoria, South Africa, in 1971. Musk is the CEO of Tesla and SpaceX. He also founded Neuralink and The Boring Company. His goal is to make life multi-planetary. SpaceX has launched reusable rockets and reduced the cost of space travel. Tesla revolutionized the electric vehicle industry with sustainable energy solutions. It also focuses on solar energy and battery storage. Neuralink is developing brain-computer interfaces to help people with neurological conditions. The Boring Company builds underground transportation tunnels to reduce urban traffic. Musk also acquired Twitter, now called X, aiming to create a free and open digital platform. He is known for his bold ideas, relentless innovation, and ambitious projects. His leadership style is characterized by risk-taking and hands-on involvement. Musk has faced criticism and challenges but continues to push boundaries. His contributions to space exploration, sustainable energy, and artificial intelligence have left a significant mark on the world. Through his companies, he aims to solve global challenges and improve the future of humanity. From launching rockets to building electric vehicles and exploring brain technology, Musk’s vision remains a driving force in technological advancement.
"""
desc = desc.split(". ")
print(desc)
upper_table.cell(4,0).merge(upper_table.cell(4,1))
para = upper_table.cell(4,0).paragraphs[0]
run = para.add_run("Specification/Description of Work (extend on to attached sheets as necessary): ")
run.bold=True
para.add_run("""The Bidder is requested to read this document carefully and accordingly submit their detailed Technical and Commercial offer and give full compliance to this Scope of Work document in their Technical and Commercial offer. """)
max_no_of_words = 600
current_no_of_words=0
for i in desc:
    if current_no_of_words<=max_no_of_words:
        current_no_of_words+=len(i.split(" "))
        para.add_run(i)
        current_no_of_words +=1
    else:
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)

        new_section.top_margin=Pt(20)
        new_section.left_margin=Pt(40)
        new_section.right_margin=Pt(40)
        new_section.bottom_margin=Pt(40)
        doc.add_picture("utils\\assets\\image_.png",width=Pt(530))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        new_sectio_tabel = doc.add_table(rows=1,cols=2)
        new_sectio_tabel.cell(0,0).merge(new_sectio_tabel.cell(0,1))

        para = new_sectio_tabel.cell(0,0).paragraphs[0]
        run = para.add_run(i)
        new_sectio_tabel.style = 'Table Grid'
        current_no_of_words = len(i.split(" "))
        max_no_of_words=800

#--------------------------------------footer section-------------------------------#
for sec in doc.sections:
    sec.footer.is_linked_to_previous=False
    sec.footer.add_paragraph(f"Issue No. :{issue_no}{' '*80}{page_no}{' '*70} issue data:{date_}")
    page_no+=1


#---------------------------------------------save section---------------------------------#


upper_table.style = 'Table Grid'

doc.save("Complex_Header_Document.docx")

print("Document with complex header created successfully!")
